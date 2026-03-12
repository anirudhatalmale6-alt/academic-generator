"""
Core DOCX document builder with native Word footnotes (OXML), TOC field codes,
cover page, and Romanian academic formatting.
"""

import re
import copy
from pathlib import Path
from typing import Optional
from lxml import etree
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.part import Part
from docx.opc.packuri import PackURI
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsmap

from core.formatting import (
    FONT_BODY, FONT_SIZE_BODY, FONT_SIZE_FOOTNOTE, FONT_SIZE_H1, FONT_SIZE_H2,
    FONT_SIZE_H3, FONT_SIZE_COVER_TITLE, FONT_SIZE_COVER_SUBTITLE, FONT_COLOR,
    LINE_SPACING_BODY, LINE_SPACING_FOOTNOTE, MARGIN_LEFT, MARGIN_RIGHT,
    MARGIN_TOP, MARGIN_BOTTOM, FIRST_LINE_INDENT, HANGING_INDENT_BIBLIO,
    ALIGN_BODY, ALIGN_CENTER, ALIGN_LEFT, HEADING_STYLES,
)

WPML_URI = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


class AcademicDocBuilder:
    """Builds a fully formatted academic DOCX with native Word features."""

    def __init__(self):
        self.doc = Document()
        self._footnote_counter = 0
        self._footnotes_part = None
        self._headings_for_toc = []  # Track headings for static TOC
        self._setup_styles()
        self._setup_page_layout()
        self._set_update_fields()

    # ─── Page layout ───────────────────────────────────────────────────

    def _setup_page_layout(self):
        """Set margins and page size (A4). Continuous footnote numbering."""
        for section in self.doc.sections:
            section.top_margin = MARGIN_TOP
            section.bottom_margin = MARGIN_BOTTOM
            section.left_margin = MARGIN_LEFT
            section.right_margin = MARGIN_RIGHT
            section.page_width = Cm(21)
            section.page_height = Cm(29.7)
            # Ensure no per-page footnote restart (continuous numbering)
            self._ensure_continuous_footnotes(section)

    def _ensure_continuous_footnotes(self, section):
        """Remove any w:numRestart from section to ensure continuous footnote numbering."""
        sectPr = section._sectPr
        for existing in sectPr.findall(qn("w:footnotePr")):
            sectPr.remove(existing)

    def _set_update_fields(self):
        """Set w:updateFields in document settings so Word auto-updates TOC on open."""
        settings_element = self.doc.settings.element
        # Remove existing updateFields if any
        for existing in settings_element.findall(qn("w:updateFields")):
            settings_element.remove(existing)
        update = OxmlElement("w:updateFields")
        update.set(qn("w:val"), "true")
        settings_element.append(update)

    # ─── Styles ────────────────────────────────────────────────────────

    def _setup_styles(self):
        """Configure document styles for Romanian academic formatting."""
        styles = self.doc.styles

        # Normal style
        normal = styles["Normal"]
        normal.font.name = FONT_BODY
        normal.font.size = FONT_SIZE_BODY
        normal.font.color.rgb = FONT_COLOR
        pf = normal.paragraph_format
        pf.line_spacing = LINE_SPACING_BODY
        pf.alignment = ALIGN_BODY
        pf.first_line_indent = FIRST_LINE_INDENT
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)

        # Set font for East Asian and Complex Script
        rpr = normal.element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.insert(0, rfonts)
        rfonts.set(qn("w:ascii"), FONT_BODY)
        rfonts.set(qn("w:hAnsi"), FONT_BODY)
        rfonts.set(qn("w:cs"), FONT_BODY)
        rfonts.set(qn("w:eastAsia"), FONT_BODY)

        # Heading styles
        for level, cfg in HEADING_STYLES.items():
            style_name = f"Heading {level}"
            style = styles[style_name]
            style.font.name = FONT_BODY
            style.font.size = cfg["size"]
            style.font.bold = cfg["bold"]
            style.font.italic = cfg["italic"]
            style.font.color.rgb = FONT_COLOR
            style.paragraph_format.alignment = cfg["alignment"]
            style.paragraph_format.first_line_indent = Cm(0)
            style.paragraph_format.space_before = Pt(12)
            style.paragraph_format.space_after = Pt(6)
            # Set font refs
            rpr = style.element.get_or_add_rPr()
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                rpr.insert(0, rfonts)
            rfonts.set(qn("w:ascii"), FONT_BODY)
            rfonts.set(qn("w:hAnsi"), FONT_BODY)
            rfonts.set(qn("w:cs"), FONT_BODY)

        # Ensure FootnoteReference character style exists with superscript
        self._ensure_footnote_reference_style()

    def _ensure_footnote_reference_style(self):
        """Create or configure the FootnoteReference character style with superscript."""
        styles = self.doc.styles
        try:
            fn_style = styles["Footnote Reference"]
        except KeyError:
            # Create the style if it doesn't exist
            fn_style = styles.add_style("Footnote Reference", 2)  # 2 = CHARACTER type

        fn_style.font.superscript = True
        fn_style.font.size = FONT_SIZE_BODY

    # ─── Cover Page ────────────────────────────────────────────────────

    def add_cover_page(
        self,
        title: str,
        doc_type: str,
        university: str = "",
        faculty: str = "",
        specialisation: str = "",
        supervisor: str = "",
        student: str = "",
        city: str = "",
        year: str = "",
    ):
        """Add a centered cover page with academic metadata."""
        # University & Faculty
        if university:
            p = self.doc.add_paragraph()
            p.alignment = ALIGN_CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            r = p.add_run(university.upper())
            r.bold = True
            r.font.name = FONT_BODY
            r.font.size = Pt(14)
            r.font.color.rgb = FONT_COLOR

        if faculty:
            p = self.doc.add_paragraph()
            p.alignment = ALIGN_CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            r = p.add_run(faculty)
            r.font.name = FONT_BODY
            r.font.size = Pt(14)
            r.font.color.rgb = FONT_COLOR

        if specialisation:
            p = self.doc.add_paragraph()
            p.alignment = ALIGN_CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            r = p.add_run(f"Specializarea: {specialisation}")
            r.font.name = FONT_BODY
            r.font.size = Pt(12)
            r.font.color.rgb = FONT_COLOR

        # Spacing before title
        for _ in range(4):
            p = self.doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0)

        # Document type label
        type_labels = {
            "licență": "LUCRARE DE LICENȚĂ",
            "disertație": "LUCRARE DE DISERTAȚIE",
            "doctorat": "TEZĂ DE DOCTORAT",
            "grad didactic": "LUCRARE PENTRU GRADUL DIDACTIC",
            "diplomă": "LUCRARE DE DIPLOMĂ",
        }
        label = type_labels.get(doc_type, doc_type.upper())
        p = self.doc.add_paragraph()
        p.alignment = ALIGN_CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        r = p.add_run(label)
        r.bold = True
        r.font.name = FONT_BODY
        r.font.size = Pt(16)
        r.font.color.rgb = FONT_COLOR

        # Title
        p = self.doc.add_paragraph()
        p.alignment = ALIGN_CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(18)
        r = p.add_run(title)
        r.bold = True
        r.font.name = FONT_BODY
        r.font.size = FONT_SIZE_COVER_TITLE
        r.font.color.rgb = FONT_COLOR

        # Spacing before metadata
        for _ in range(4):
            p = self.doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0)

        # Supervisor and student
        if supervisor or student:
            p = self.doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0)
            if supervisor:
                r = p.add_run(f"Coordonator științific:\n{supervisor}")
                r.font.name = FONT_BODY
                r.font.size = Pt(12)
                r.font.color.rgb = FONT_COLOR

        if student:
            p = self.doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0)
            p.alignment = ALIGN_LEFT
            r = p.add_run(f"Absolvent:\n{student}")
            r.font.name = FONT_BODY
            r.font.size = Pt(12)
            r.font.color.rgb = FONT_COLOR

        # Spacing
        for _ in range(3):
            p = self.doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0)

        # City and year
        footer_text = ""
        if city:
            footer_text += city
        if year:
            if footer_text:
                footer_text += ", "
            footer_text += str(year)
        if footer_text:
            p = self.doc.add_paragraph()
            p.alignment = ALIGN_CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            r = p.add_run(footer_text)
            r.font.name = FONT_BODY
            r.font.size = Pt(14)
            r.font.color.rgb = FONT_COLOR

        # Page break after cover — no page number on cover
        self._add_section_break_no_page_number()

    def _add_section_break_no_page_number(self):
        """Add a section break and suppress page number on current section."""
        new_section = self.doc.add_section()
        new_section.top_margin = MARGIN_TOP
        new_section.bottom_margin = MARGIN_BOTTOM
        new_section.left_margin = MARGIN_LEFT
        new_section.right_margin = MARGIN_RIGHT
        self._ensure_continuous_footnotes(new_section)

    # ─── Abstract / Summary ───────────────────────────────────────────

    def add_abstract(self, text: str, title: str = "Rezumat"):
        """Add abstract/summary page."""
        self.doc.add_heading(title, level=1)
        self._headings_for_toc.append({"title": title, "level": 1})
        for para in text.strip().split("\n\n"):
            para = para.strip()
            if para:
                p = self.doc.add_paragraph(para)
                self._apply_body_format(p)
        self.doc.add_page_break()

    # ─── Table of Contents ─────────────────────────────────────────────

    def add_toc(self, title: str = "Cuprins", sections: list[dict] | None = None):
        """Add TOC with both a field code (auto-updates in Word) and static entries.

        The w:updateFields setting in document.xml tells Word to refresh the TOC
        automatically when the document is opened. Static entries are also written
        as fallback content inside the field.

        Args:
            sections: List of dicts with 'title' and 'level' keys for static TOC entries.
        """
        # TOC heading
        self.doc.add_heading(title, level=1)

        # Build TOC field code paragraph
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)

        # --- Field BEGIN ---
        r_begin = OxmlElement("w:r")
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        r_begin.append(fld_begin)
        p._p.append(r_begin)

        # --- Field INSTRUCTION ---
        r_instr = OxmlElement("w:r")
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = r" TOC \o '1-3' \h \z \u "
        r_instr.append(instr_text)
        p._p.append(r_instr)

        # --- Field SEPARATE ---
        r_sep = OxmlElement("w:r")
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        r_sep.append(fld_sep)
        p._p.append(r_sep)

        # --- Static TOC entries (fallback content between SEPARATE and END) ---
        toc_entries = sections if sections else self._headings_for_toc
        if toc_entries:
            for entry in toc_entries:
                toc_level = entry.get("level", 1)
                toc_title = entry.get("title", "")
                indent = "    " * (toc_level - 1)
                toc_p = self.doc.add_paragraph()
                toc_p.paragraph_format.first_line_indent = Cm(0)
                toc_p.paragraph_format.space_before = Pt(2)
                toc_p.paragraph_format.space_after = Pt(2)
                if toc_level > 1:
                    toc_p.paragraph_format.left_indent = Cm(1.0 * (toc_level - 1))
                r = toc_p.add_run(toc_title)
                r.font.name = FONT_BODY
                r.font.size = FONT_SIZE_BODY
                r.font.color.rgb = FONT_COLOR
                if toc_level == 1:
                    r.bold = True
        else:
            # Minimal placeholder if no sections known
            r_placeholder = OxmlElement("w:r")
            t_placeholder = OxmlElement("w:t")
            t_placeholder.text = "[Cuprinsul se va actualiza automat la deschiderea documentului în Word]"
            r_placeholder.append(t_placeholder)
            p._p.append(r_placeholder)

        # --- Field END --- (goes in a new paragraph after TOC entries)
        p_end = self.doc.add_paragraph()
        p_end.paragraph_format.first_line_indent = Cm(0)
        r_end = OxmlElement("w:r")
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        r_end.append(fld_end)
        p_end._p.append(r_end)

        self.doc.add_page_break()

    def track_heading(self, title: str, level: int):
        """Track a heading for static TOC generation."""
        self._headings_for_toc.append({"title": title, "level": level})

    # ─── Page Numbers ──────────────────────────────────────────────────

    def add_page_numbers(self):
        """Add bottom-center page numbers starting from the current section."""
        for section in self.doc.sections:
            footer = section.footer
            footer.is_linked_to_previous = False
            p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            p.alignment = ALIGN_CENTER
            p.paragraph_format.first_line_indent = Cm(0)

            # PAGE field code
            run = p.add_run()
            fld_begin = OxmlElement("w:fldChar")
            fld_begin.set(qn("w:fldCharType"), "begin")
            run._r.append(fld_begin)

            instr = OxmlElement("w:instrText")
            instr.set(qn("xml:space"), "preserve")
            instr.text = " PAGE "
            run._r.append(instr)

            fld_sep = OxmlElement("w:fldChar")
            fld_sep.set(qn("w:fldCharType"), "separate")
            run._r.append(fld_sep)

            num_run = OxmlElement("w:t")
            num_run.text = "1"
            run._r.append(num_run)

            fld_end = OxmlElement("w:fldChar")
            fld_end.set(qn("w:fldCharType"), "end")
            run._r.append(fld_end)

    # ─── Chapters / Content Sections ───────────────────────────────────

    def add_chapter(self, title: str, content: str, level: int = 1):
        """Add a chapter with heading and body paragraphs."""
        self.doc.add_heading(title, level=level)
        self._headings_for_toc.append({"title": title, "level": level})
        self._add_content_paragraphs(content)

    def add_subchapter(self, title: str, content: str, level: int = 2):
        """Add a subchapter."""
        self.add_chapter(title, content, level=level)

    def _add_content_paragraphs(self, content: str):
        """Parse content text and add formatted paragraphs.
        Handles heading markers (# H1, ## H2, ### H3) in AI output."""
        lines = content.strip().split("\n")
        current_para = []

        for line in lines:
            stripped = line.strip()
            if not stripped:
                # Flush current paragraph
                if current_para:
                    text = " ".join(current_para)
                    p = self.doc.add_paragraph(text)
                    self._apply_body_format(p)
                    current_para = []
                continue

            # Check for heading markers from AI output
            heading_match = re.match(r'^(#{1,3})\s+(.+)$', stripped)
            if heading_match:
                # Flush any pending paragraph
                if current_para:
                    text = " ".join(current_para)
                    p = self.doc.add_paragraph(text)
                    self._apply_body_format(p)
                    current_para = []
                level = len(heading_match.group(1))
                heading_text = heading_match.group(2).strip()
                self.doc.add_heading(heading_text, level=min(level, 3))
                self._headings_for_toc.append({"title": heading_text, "level": min(level, 3)})
                continue

            # Remove markdown artifacts
            stripped = re.sub(r'\*\*(.+?)\*\*', r'\1', stripped)  # Bold
            stripped = re.sub(r'\*(.+?)\*', r'\1', stripped)      # Italic
            stripped = re.sub(r'^[-•]\s+', '', stripped)           # Bullet points → normal text

            current_para.append(stripped)

        # Flush remaining
        if current_para:
            text = " ".join(current_para)
            p = self.doc.add_paragraph(text)
            self._apply_body_format(p)

    def _apply_body_format(self, paragraph):
        """Apply standard body formatting to a paragraph."""
        paragraph.paragraph_format.line_spacing = LINE_SPACING_BODY
        paragraph.paragraph_format.alignment = ALIGN_BODY
        paragraph.paragraph_format.first_line_indent = FIRST_LINE_INDENT
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.font.name = FONT_BODY
            run.font.size = FONT_SIZE_BODY
            run.font.color.rgb = FONT_COLOR

    # ─── Native Word Footnotes (OXML) ─────────────────────────────────

    def _get_or_create_footnotes_part(self):
        """Get or create the footnotes.xml part with proper OXML structure."""
        if self._footnotes_part is not None:
            return self._footnotes_part

        main_part = self.doc.part

        # Check if footnotes part already exists
        for rel in main_part.rels.values():
            if "footnotes" in rel.reltype:
                self._footnotes_part = rel.target_part
                return self._footnotes_part

        # Create footnotes.xml from scratch
        footnotes_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
            '<w:footnote w:type="separator" w:id="-1">'
            '<w:p><w:r><w:separator/></w:r></w:p>'
            '</w:footnote>'
            '<w:footnote w:type="continuationSeparator" w:id="0">'
            '<w:p><w:r><w:continuationSeparator/></w:r></w:p>'
            '</w:footnote>'
            '</w:footnotes>'
        )

        footnotes_part = Part(
            PackURI("/word/footnotes.xml"),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml",
            footnotes_xml.encode("utf-8"),
            main_part.package,
        )

        main_part.relate_to(
            footnotes_part,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes",
        )

        self._footnotes_part = footnotes_part
        return footnotes_part

    def add_footnote(self, paragraph, footnote_text: str) -> int:
        """Add a native Word footnote to a paragraph.

        Creates a w:footnoteReference in the paragraph (superscript number)
        and a w:footnote element in footnotes.xml with proper formatting.

        Returns the footnote ID.
        """
        footnotes_part = self._get_or_create_footnotes_part()
        self._footnote_counter += 1
        fn_id = self._footnote_counter

        # Parse the current footnotes XML
        footnotes_element = etree.fromstring(footnotes_part.blob)

        # Create the footnote element in footnotes.xml
        fn = etree.SubElement(footnotes_element, qn("w:footnote"))
        fn.set(qn("w:id"), str(fn_id))

        # Footnote paragraph
        fn_p = etree.SubElement(fn, qn("w:p"))

        # Paragraph properties — single spacing
        fn_ppr = etree.SubElement(fn_p, qn("w:pPr"))
        fn_spacing = etree.SubElement(fn_ppr, qn("w:spacing"))
        fn_spacing.set(qn("w:line"), "240")
        fn_spacing.set(qn("w:lineRule"), "auto")
        # Footnote text style
        fn_pstyle = etree.SubElement(fn_ppr, qn("w:pStyle"))
        fn_pstyle.set(qn("w:val"), "FootnoteText")

        # Run 1: Footnote reference mark (auto-number in footnote area)
        fn_r1 = etree.SubElement(fn_p, qn("w:r"))
        fn_r1_rpr = etree.SubElement(fn_r1, qn("w:rPr"))
        fn_r1_style = etree.SubElement(fn_r1_rpr, qn("w:rStyle"))
        fn_r1_style.set(qn("w:val"), "FootnoteReference")
        fn_r1_valign = etree.SubElement(fn_r1_rpr, qn("w:vertAlign"))
        fn_r1_valign.set(qn("w:val"), "superscript")
        fn_r1_sz = etree.SubElement(fn_r1_rpr, qn("w:sz"))
        fn_r1_sz.set(qn("w:val"), "20")  # 10pt
        fn_r1_szcs = etree.SubElement(fn_r1_rpr, qn("w:szCs"))
        fn_r1_szcs.set(qn("w:val"), "20")
        fn_r1_fonts = etree.SubElement(fn_r1_rpr, qn("w:rFonts"))
        fn_r1_fonts.set(qn("w:ascii"), FONT_BODY)
        fn_r1_fonts.set(qn("w:hAnsi"), FONT_BODY)
        fn_r1_fonts.set(qn("w:cs"), FONT_BODY)
        etree.SubElement(fn_r1, qn("w:footnoteRef"))

        # Run 2: Space
        fn_r2 = etree.SubElement(fn_p, qn("w:r"))
        fn_r2_rpr = etree.SubElement(fn_r2, qn("w:rPr"))
        fn_r2_sz = etree.SubElement(fn_r2_rpr, qn("w:sz"))
        fn_r2_sz.set(qn("w:val"), "20")
        fn_r2_szcs = etree.SubElement(fn_r2_rpr, qn("w:szCs"))
        fn_r2_szcs.set(qn("w:val"), "20")
        fn_r2_fonts = etree.SubElement(fn_r2_rpr, qn("w:rFonts"))
        fn_r2_fonts.set(qn("w:ascii"), FONT_BODY)
        fn_r2_fonts.set(qn("w:hAnsi"), FONT_BODY)
        fn_r2_t = etree.SubElement(fn_r2, qn("w:t"))
        fn_r2_t.set(qn("xml:space"), "preserve")
        fn_r2_t.text = " "

        # Run 3: Footnote text (10pt TNR)
        fn_r3 = etree.SubElement(fn_p, qn("w:r"))
        fn_r3_rpr = etree.SubElement(fn_r3, qn("w:rPr"))
        fn_r3_sz = etree.SubElement(fn_r3_rpr, qn("w:sz"))
        fn_r3_sz.set(qn("w:val"), "20")
        fn_r3_szcs = etree.SubElement(fn_r3_rpr, qn("w:szCs"))
        fn_r3_szcs.set(qn("w:val"), "20")
        fn_r3_fonts = etree.SubElement(fn_r3_rpr, qn("w:rFonts"))
        fn_r3_fonts.set(qn("w:ascii"), FONT_BODY)
        fn_r3_fonts.set(qn("w:hAnsi"), FONT_BODY)
        fn_r3_fonts.set(qn("w:cs"), FONT_BODY)
        fn_r3_t = etree.SubElement(fn_r3, qn("w:t"))
        fn_r3_t.set(qn("xml:space"), "preserve")
        fn_r3_t.text = footnote_text

        # Save updated footnotes XML
        footnotes_part._blob = etree.tostring(footnotes_element, xml_declaration=True, encoding="UTF-8", standalone=True)

        # ─── Add superscript footnote reference in the body text ───────
        p_element = paragraph._p

        # Create run with footnoteReference + explicit superscript
        ref_run = OxmlElement("w:r")
        ref_rpr = OxmlElement("w:rPr")
        # Style reference
        ref_style = OxmlElement("w:rStyle")
        ref_style.set(qn("w:val"), "FootnoteReference")
        ref_rpr.append(ref_style)
        # Explicit superscript (in case style isn't recognized)
        ref_valign = OxmlElement("w:vertAlign")
        ref_valign.set(qn("w:val"), "superscript")
        ref_rpr.append(ref_valign)
        ref_run.append(ref_rpr)
        # The actual footnote reference element
        ref_fn = OxmlElement("w:footnoteReference")
        ref_fn.set(qn("w:id"), str(fn_id))
        ref_run.append(ref_fn)
        p_element.append(ref_run)

        return fn_id

    # ─── Bibliography ──────────────────────────────────────────────────

    def add_bibliography(self, entries: list[str], title: str = "Bibliografie"):
        """Add bibliography section with hanging indent formatting."""
        self.doc.add_page_break()
        self.doc.add_heading(title, level=1)
        self._headings_for_toc.append({"title": title, "level": 1})

        # Sort entries alphabetically
        sorted_entries = sorted(entries, key=lambda x: x.strip().lower())

        for entry in sorted_entries:
            entry = entry.strip()
            if not entry:
                continue
            p = self.doc.add_paragraph(entry)
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.alignment = ALIGN_BODY
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.left_indent = HANGING_INDENT_BIBLIO
            # Hanging indent via negative first line
            ppr = p._p.get_or_add_pPr()
            ind = ppr.find(qn("w:ind"))
            if ind is None:
                ind = OxmlElement("w:ind")
                ppr.append(ind)
            ind.set(qn("w:left"), str(int(HANGING_INDENT_BIBLIO.emu / 635)))
            ind.set(qn("w:hanging"), str(int(HANGING_INDENT_BIBLIO.emu / 635)))

            for run in p.runs:
                run.font.name = FONT_BODY
                run.font.size = FONT_SIZE_BODY
                run.font.color.rgb = FONT_COLOR

    # ─── Save ──────────────────────────────────────────────────────────

    def save(self, filepath: Path):
        """Save the document."""
        filepath = Path(filepath)
        filepath.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(filepath))
        return filepath
