"""
Editor page — Modify an existing document by replacing only the requested sections.
Preserves the original document's formatting, cover page, TOC, footnotes, and bibliography.
"""

import re
import copy
import streamlit as st
from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from core.registry import DocumentRegistry
from core.ai_providers import generate_text, get_available_providers
from core.prompt_builder import build_system_prompt
from core.guide_reader import extract_text as extract_guide_text


def _parse_doc_sections(doc):
    """Parse a Document into sections grouped by heading.

    Returns list of dicts: {heading, heading_idx, level, paragraphs: [(idx, text)]}
    """
    sections = []
    current = None

    for i, para in enumerate(doc.paragraphs):
        if para.style.name.startswith("Heading"):
            try:
                level = int(para.style.name.split()[-1])
            except ValueError:
                level = 1
            if current is not None:
                sections.append(current)
            current = {
                "heading": para.text.strip(),
                "heading_idx": i,
                "level": level,
                "paragraphs": [],
            }
        elif current is not None and para.text.strip():
            current["paragraphs"].append((i, para.text.strip()))

    if current is not None:
        sections.append(current)

    return sections


def _parse_ai_modifications(ai_output: str) -> dict[str, str]:
    """Parse AI output for modified sections.

    Expected format:
    === MODIFIED: Section Title ===
    content...
    === END ===
    """
    modifications = {}
    pattern = re.compile(
        r'===\s*MODIFIED:\s*(.+?)\s*===\s*\n(.*?)\n\s*===\s*END\s*===',
        re.DOTALL | re.IGNORECASE,
    )
    for match in pattern.finditer(ai_output):
        title = match.group(1).strip()
        content = match.group(2).strip()
        if title and content:
            modifications[title] = content

    return modifications


def _replace_section_in_doc(doc, heading_title: str, new_text: str) -> bool:
    """Replace the content of a section in-place, preserving the rest of the document.

    Finds the heading, removes all paragraphs between it and the next heading,
    then inserts new formatted paragraphs.
    Returns True if the section was found and modified.
    """
    body = doc.element.body
    all_paras = list(doc.paragraphs)

    # Find heading paragraph (case-insensitive match)
    heading_idx = None
    for i, p in enumerate(all_paras):
        if p.style.name.startswith("Heading") and p.text.strip().lower() == heading_title.strip().lower():
            heading_idx = i
            break

    if heading_idx is None:
        # Try partial match (heading might have been slightly different)
        target = heading_title.strip().lower()
        for i, p in enumerate(all_paras):
            if p.style.name.startswith("Heading") and target in p.text.strip().lower():
                heading_idx = i
                break

    if heading_idx is None:
        return False

    # Find the end of this section (next heading or end of document)
    end_idx = len(all_paras)
    for i in range(heading_idx + 1, len(all_paras)):
        if all_paras[i].style.name.startswith("Heading"):
            end_idx = i
            break

    # Collect elements to remove (content paragraphs only, not the heading)
    elements_to_remove = []
    for i in range(heading_idx + 1, end_idx):
        elements_to_remove.append(all_paras[i]._element)

    # Remove content paragraphs
    for elem in elements_to_remove:
        body.remove(elem)

    # Parse new text into paragraphs
    new_paras_text = [p.strip() for p in new_text.strip().split('\n\n') if p.strip()]
    # If only single newlines used, split by those
    if len(new_paras_text) <= 1 and '\n' in new_text:
        new_paras_text = [p.strip() for p in new_text.strip().split('\n') if p.strip()]

    # Insert new paragraphs after the heading element
    insert_after = all_paras[heading_idx]._element
    for text in new_paras_text:
        # Strip markdown formatting
        text = re.sub(r'\*{1,3}([^*]+)\*{1,3}', r'\1', text)
        text = re.sub(r'^[-•]\s+', '', text, flags=re.MULTILINE)

        # Check if it's a sub-heading
        heading_match = re.match(r'^(#{1,3})\s+(.+)$', text)
        if heading_match:
            level = min(len(heading_match.group(1)), 3)
            h_text = heading_match.group(2).strip()
            new_p_elem = _create_heading_element(h_text, level)
        else:
            new_p_elem = _create_body_paragraph(text)

        insert_after.addnext(new_p_elem)
        insert_after = new_p_elem

    return True


def _create_body_paragraph(text: str):
    """Create a properly formatted body paragraph element."""
    p = OxmlElement("w:p")

    # Paragraph properties: 1.5 spacing, justified, first-line indent
    pPr = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:line"), "360")  # 1.5 line spacing
    spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "both")
    pPr.append(jc)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:firstLine"), "709")  # ~1.25cm first-line indent
    pPr.append(ind)
    p.append(pPr)

    # Run with TNR 12pt
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Times New Roman")
    fonts.set(qn("w:hAnsi"), "Times New Roman")
    fonts.set(qn("w:cs"), "Times New Roman")
    rPr.append(fonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "24")  # 12pt
    rPr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), "24")
    rPr.append(szCs)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    p.append(r)

    return p


def _create_heading_element(text: str, level: int):
    """Create a heading paragraph element."""
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    pStyle = OxmlElement("w:pStyle")
    pStyle.set(qn("w:val"), f"Heading{level}")
    pPr.append(pStyle)
    p.append(pPr)

    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Times New Roman")
    fonts.set(qn("w:hAnsi"), "Times New Roman")
    fonts.set(qn("w:cs"), "Times New Roman")
    rPr.append(fonts)
    if level == 1:
        b = OxmlElement("w:b")
        rPr.append(b)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    p.append(r)

    return p


def render(app_dir: Path):
    st.title("Editor Document")
    st.markdown(
        "Încărcați un document și descrieți modificările dorite. "
        "Editorul va modifica **doar** secțiunile solicitate, păstrând restul documentului intact."
    )

    registry_path = app_dir / "registry.json"
    output_dir = app_dir / "outputs"
    registry = DocumentRegistry(registry_path)

    # ─── Source selection ──────────────────────────────────────────
    source = st.radio("Sursă document", ["Din registru", "Încărcare fișier"], horizontal=True)

    doc_path = None
    doc_code = ""

    if source == "Din registru":
        entries = registry.list_all()
        if not entries:
            st.info("Registrul este gol. Generați un document mai întâi.")
            return

        options = {f"{e['code']} — {e['title']}": e for e in entries}
        selected = st.selectbox("Selectați documentul", list(options.keys()))
        if selected:
            entry = options[selected]
            doc_path = Path(entry["output_path"])
            doc_code = entry["code"]
            if not doc_path.exists():
                st.error(f"Fișierul nu a fost găsit: {doc_path}")
                return
    else:
        uploaded = st.file_uploader("Încărcați document .docx", type=["docx"])
        if uploaded:
            tmp_path = output_dir / "_editor_temp" / uploaded.name
            tmp_path.parent.mkdir(parents=True, exist_ok=True)
            tmp_path.write_bytes(uploaded.read())
            doc_path = tmp_path
            doc_code = tmp_path.stem

    if doc_path is None:
        return

    # ─── Load and display document structure ───────────────────────
    try:
        doc = Document(str(doc_path))
    except Exception as e:
        st.error(f"Eroare la deschiderea documentului: {e}")
        return

    sections = _parse_doc_sections(doc)

    st.markdown("---")
    st.subheader("Structura documentului")

    if sections:
        for sec in sections:
            indent = "  " * (sec["level"] - 1)
            para_count = len(sec["paragraphs"])
            st.markdown(f"{indent}**{sec['heading']}** ({para_count} paragrafe)")
    else:
        st.info("Nu s-au găsit titluri de secțiuni în document.")

    # ─── Modification instructions ─────────────────────────────────
    st.markdown("---")
    st.subheader("Instrucțiuni de Modificare")

    modification = st.text_area(
        "Descrieți toate modificările dorite",
        height=200,
        help="Descrieți liber ce modificări doriți. Menționați secțiunile specifice pe care doriți să le modificați.",
        placeholder="Ex: Rescrie secțiunea Concluzii cu accent pe rezultatele practice. "
                    "Adaugă mai multe detalii în Introducere despre contextul economic actual.",
    )

    # ─── Reference file uploads ────────────────────────────────────
    ref_files = st.file_uploader(
        "Fișiere de referință (opțional)",
        type=["pdf", "docx", "txt", "csv", "xlsx", "md"],
        accept_multiple_files=True,
        help="Fișiere pe care AI-ul ar trebui să le ia în considerare la modificare.",
    )

    providers = get_available_providers()
    provider_list = list(providers.keys())

    if st.button("Aplică Modificările", type="primary", disabled=not modification.strip(), use_container_width=True):
        if not provider_list:
            st.error("Niciun provider AI configurat!")
            return

        if not sections:
            st.error("Documentul nu conține secțiuni cu titluri.")
            return

        # Build section text for AI context
        section_text = ""
        for sec in sections:
            section_text += f"\n=== {sec['heading']} ===\n"
            for _, para_text in sec["paragraphs"]:
                section_text += para_text + "\n\n"

        # Truncate if too long (keep first 15000 chars for context)
        if len(section_text) > 15000:
            section_text = section_text[:15000] + "\n\n[... conținut trunchiat ...]"

        # Extract reference file content
        ref_context = ""
        if ref_files:
            ref_dir = output_dir / "_editor_temp" / "refs"
            ref_dir.mkdir(parents=True, exist_ok=True)
            for rf in ref_files:
                ref_path = ref_dir / rf.name
                ref_path.write_bytes(rf.read())
                text = extract_guide_text(ref_path)
                if text:
                    ref_context += f"\n--- Fișier referință: {rf.name} ---\n{text[:3000]}\n"

        with st.spinner("Se aplică modificările..."):
            prompt = f"""Ai mai jos conținutul unui document academic organizat pe secțiuni:

{section_text}
{f'''
--- FIȘIERE DE REFERINȚĂ ---
{ref_context}
--- SFÂRȘIT REFERINȚE ---
''' if ref_context else ''}
Instrucțiuni de modificare de la utilizator:
{modification}

REGULI IMPORTANTE:
1. Returnează DOAR secțiunile pe care le-ai modificat — NU returna secțiunile nemodificate
2. Folosește EXACT acest format pentru fiecare secțiune modificată:

=== MODIFIED: Titlul Exact al Secțiunii ===
[Conținutul modificat - text simplu, fără markdown]
=== END ===

3. Titlul secțiunii trebuie să fie EXACT ca în documentul original
4. Păstrează stilul academic și tonul formal
5. Păstrează citările existente în format (Autor, An)
6. NU folosi markdown bold/italic (*text*)
7. Scrie paragrafe complete separate prin linii goale
8. Conținutul modificat trebuie să fie cel puțin la fel de lung ca originalul"""

            try:
                new_content, provider_used = generate_text(
                    prompt=prompt,
                    system_prompt=build_system_prompt(provider_list[0]),
                    provider_order=provider_list,
                    max_tokens=8192,
                )

                # Parse modified sections from AI output
                modifications = _parse_ai_modifications(new_content)

                if not modifications:
                    st.warning(
                        "AI-ul nu a returnat modificări în formatul așteptat. "
                        "Încercați să reformulați instrucțiunile."
                    )
                    with st.expander("Răspuns AI (debug)"):
                        st.text_area("Output", value=new_content, height=300, disabled=True)
                    return

                # Apply modifications in-place on the loaded document
                applied = []
                not_found = []
                for sec_title, sec_content in modifications.items():
                    success = _replace_section_in_doc(doc, sec_title, sec_content)
                    if success:
                        applied.append(sec_title)
                    else:
                        not_found.append(sec_title)

                if not applied:
                    st.error("Nu s-a putut aplica nicio modificare — titlurile secțiunilor nu au fost găsite în document.")
                    with st.expander("Secțiuni returnate de AI"):
                        for title in not_found:
                            st.write(f"- {title}")
                    return

                # Save as new version
                stem = doc_path.stem
                version_match = re.search(r'_v(\d+)$', stem)
                if version_match:
                    version = int(version_match.group(1)) + 1
                    new_stem = re.sub(r'_v\d+$', f'_v{version}', stem)
                else:
                    new_stem = f"{stem}_v2"

                new_path = doc_path.parent / f"{new_stem}.docx"
                doc.save(str(new_path))

                st.success(
                    f"Modificări aplicate cu succes folosind {provider_used}!\n\n"
                    f"Secțiuni modificate: {', '.join(applied)}"
                )

                if not_found:
                    st.warning(f"Secțiuni negăsite: {', '.join(not_found)}")

                with open(new_path, "rb") as f:
                    st.download_button(
                        "Descarcă documentul modificat",
                        data=f.read(),
                        file_name=new_path.name,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )

            except Exception as e:
                st.error(f"Eroare: {e}")
                import traceback
                st.expander("Detalii eroare").code(traceback.format_exc())
